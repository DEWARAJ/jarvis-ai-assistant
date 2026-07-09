"""Real file-system operations for JARVIS — the 'doer' tool.

Safety model (matches SAFETY_PROTOCOL):
  - System/protected directories are blocked outright.
  - Read-only actions (scan, find, read) run freely.
  - Changes (organize/move/rename/delete) are PLANNED and previewed first; the
    orchestrator only executes them after explicit Master confirmation.
  - 'Delete' never hard-deletes: files are moved to a `_jarvis_recovery` folder.
Pure standard library.
"""
from __future__ import annotations
import os, shutil, fnmatch
from datetime import datetime
from tools.base_tool import BaseTool

CATEGORIES = {
    "Images": {"jpg","jpeg","png","gif","bmp","webp","svg","heic","tiff","ico"},
    "Documents": {"pdf","doc","docx","txt","md","rtf","odt","xls","xlsx","csv","ppt","pptx","epub"},
    "Videos": {"mp4","mov","avi","mkv","webm","flv","wmv"},
    "Audio": {"mp3","wav","flac","aac","ogg","m4a"},
    "Archives": {"zip","rar","7z","tar","gz","bz2"},
    "Installers": {"exe","msi","dmg","appimage"},
    "Code": {"py","js","ts","html","css","json","java","c","cpp","cs","go","rs","rb","php","sh","sql"},
}
READABLE_TEXT = {"txt","md","csv","json","log","py","js","ts","html","css","ini","yaml","yml","sh","sql"}
RECOVERY_DIR = "_jarvis_recovery"


class _MissingLib(Exception):
    """Raised when an optional document-reading library isn't installed."""

# folder shortcuts -> real user paths
SHORTCUTS = {
    "downloads": "~/Downloads", "download": "~/Downloads",
    "desktop": "~/Desktop", "documents": "~/Documents", "docs": "~/Documents",
    "pictures": "~/Pictures", "music": "~/Music", "videos": "~/Videos",
}


def category_of(ext: str) -> str:
    ext = ext.lower().lstrip(".")
    for cat, exts in CATEGORIES.items():
        if ext in exts:
            return cat
    return "Other"


class FileOpsTool(BaseTool):
    name = "file_ops"; scope = "user folders (preview + confirm)"

    # ---- path handling & safety ----
    def resolve(self, raw: str) -> str:
        raw = (raw or "").strip().strip('"').strip("'")
        low = raw.lower()
        if low in SHORTCUTS:
            raw = SHORTCUTS[low]
        return os.path.abspath(os.path.expanduser(raw))

    def is_protected(self, path: str) -> bool:
        p = path.replace("\\", "/").lower().rstrip("/")
        # drive roots like c:/ or filesystem root
        if len(p) <= 3:
            return True
        blocked = ["/windows", "/program files", "/programdata", "/system32",
                   "/$recycle", "/boot", "/etc", "/usr/", "/bin", "/sbin",
                   "/library/", "/system/"]
        return any(b in p for b in blocked)

    # ---- read-only ----
    def scan(self, raw: str) -> str:
        path = self.resolve(raw)
        if not os.path.isdir(path):
            return f"I can't find a folder at: {path}"
        if self.is_protected(path):
            return f"That looks like a protected system folder — I won't touch {path}."
        try:
            entries = os.listdir(path)
        except OSError as e:
            return f"Can't read that folder ({e})."
        files = [e for e in entries if os.path.isfile(os.path.join(path, e))]
        folders = [e for e in entries if os.path.isdir(os.path.join(path, e))]
        by_cat, total = {}, 0
        for f in files:
            ext = os.path.splitext(f)[1]
            cat = category_of(ext)
            by_cat[cat] = by_cat.get(cat, 0) + 1
            try: total += os.path.getsize(os.path.join(path, f))
            except OSError: pass
        cats = ", ".join(f"{k}: {v}" for k, v in sorted(by_cat.items())) or "nothing"
        return (f"Folder: {path}\n"
                f"{len(files)} files, {len(folders)} subfolders, {total/1e6:.1f} MB total.\n"
                f"By type — {cats}.")

    def find(self, raw: str, pattern: str) -> str:
        path = self.resolve(raw)
        if not os.path.isdir(path):
            return f"I can't find a folder at: {path}"
        pat = pattern if any(c in pattern for c in "*?.") else f"*{pattern}*"
        hits = []
        for root, _dirs, files in os.walk(path):
            for f in files:
                if fnmatch.fnmatch(f.lower(), pat.lower()):
                    hits.append(os.path.join(root, f))
            if len(hits) > 200:
                break
        if not hits:
            return f"No files matching '{pattern}' under {path}."
        head = "\n".join(f"  {h}" for h in hits[:25])
        more = f"\n  ...and {len(hits)-25} more" if len(hits) > 25 else ""
        return f"Found {len(hits)} match(es) for '{pattern}':\n{head}{more}"

    def read_text(self, raw: str, max_chars: int = 12000) -> tuple[str, str]:
        """Return (status, content). Handles text/code, PDF, Word (.docx), Excel (.xlsx).
        content is empty on failure (caller falls back / shows the status)."""
        path = self.resolve(raw)
        if not os.path.isfile(path):
            return (f"I can't find a file at: {path}", "")
        ext = os.path.splitext(path)[1].lower().lstrip(".")
        try:
            if ext == "pdf":
                data = self._read_pdf(path, max_chars)
            elif ext == "docx":
                data = self._read_docx(path, max_chars)
            elif ext in ("xlsx", "xlsm"):
                data = self._read_xlsx(path, max_chars)
            elif ext in READABLE_TEXT:
                with open(path, encoding="utf-8", errors="replace") as f:
                    data = f.read(max_chars)
            elif ext == "doc":
                return ("That's an old-format .doc file. Please re-save it as .docx and I'll read it.", "")
            else:
                return (f"That's a .{ext} file — I can read text, code, PDF, Word (.docx) and Excel (.xlsx). "
                        f"For this type, try converting it first.", "")
        except _MissingLib as e:
            return (str(e), "")
        except Exception as e:
            return (f"I found the file but couldn't read its contents ({e}).", "")
        if not (data and data.strip()):
            return (f"I opened {os.path.basename(path)} but it looks empty or has no extractable text "
                    "(scanned PDFs need OCR, which is a later upgrade).", "")
        return (f"Read {os.path.basename(path)} ({len(data)} chars extracted).", data)

    def _read_pdf(self, path: str, max_chars: int) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader  # fallback name
            except ImportError:
                raise _MissingLib("To read PDFs, install the reader once: "
                                  "pip install pypdf  (then restart JARVIS).")
        reader = PdfReader(path)
        out = []
        for page in reader.pages:
            out.append(page.extract_text() or "")
            if sum(len(x) for x in out) > max_chars:
                break
        return "\n".join(out)[:max_chars]

    def _read_docx(self, path: str, max_chars: int) -> str:
        try:
            import docx  # python-docx
        except ImportError:
            raise _MissingLib("To read Word files, install it once: "
                              "pip install python-docx  (then restart JARVIS).")
        d = docx.Document(path)
        parts = [para.text for para in d.paragraphs if para.text]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)[:max_chars]

    def _read_xlsx(self, path: str, max_chars: int) -> str:
        try:
            import openpyxl
        except ImportError:
            raise _MissingLib("To read Excel files, install it once: "
                              "pip install openpyxl  (then restart JARVIS).")
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    lines.append(" | ".join(cells))
                if sum(len(x) for x in lines) > max_chars:
                    break
        wb.close()
        return "\n".join(lines)[:max_chars]

    # ---- change: plan then execute ----
    def plan_organize(self, raw: str) -> dict:
        path = self.resolve(raw)
        if not os.path.isdir(path):
            return {"ok": False, "msg": f"I can't find a folder at: {path}"}
        if self.is_protected(path):
            return {"ok": False, "msg": f"That's a protected system folder — I won't reorganize {path}."}
        moves = []
        for e in os.listdir(path):
            full = os.path.join(path, e)
            if not os.path.isfile(full):
                continue
            cat = category_of(os.path.splitext(e)[1])
            moves.append((e, cat))
        if not moves:
            return {"ok": False, "msg": f"No loose files to organize in {path}."}
        counts = {}
        for _, c in moves:
            counts[c] = counts.get(c, 0) + 1
        preview = ", ".join(f"{v} -> {k}/" for k, v in sorted(counts.items()))
        return {"ok": True, "path": path, "moves": moves,
                "msg": f"Plan for {path}:\nMove {len(moves)} files into subfolders ({preview}).",
                "summary": counts}

    def execute_organize(self, plan: dict) -> str:
        path = plan["path"]; done = 0; failed = 0
        for name, cat in plan["moves"]:
            src = os.path.join(path, name)
            dst_dir = os.path.join(path, cat)
            try:
                os.makedirs(dst_dir, exist_ok=True)
                dst = os.path.join(dst_dir, name)
                if os.path.abspath(src) != os.path.abspath(dst):
                    shutil.move(src, dst)
                done += 1
            except OSError:
                failed += 1
        msg = f"Done. Organized {done} files into subfolders in {path}."
        if failed:
            msg += f" ({failed} couldn't be moved and were left in place.)"
        return msg

    def plan_delete(self, raw: str, pattern: str) -> dict:
        path = self.resolve(raw)
        if not os.path.isdir(path):
            return {"ok": False, "msg": f"I can't find a folder at: {path}"}
        if self.is_protected(path):
            return {"ok": False, "msg": f"Protected system folder — I won't delete from {path}."}
        pat = pattern if any(c in pattern for c in "*?.") else f"*{pattern}*"
        targets = [e for e in os.listdir(path)
                   if os.path.isfile(os.path.join(path, e)) and fnmatch.fnmatch(e.lower(), pat.lower())]
        if not targets:
            return {"ok": False, "msg": f"No files matching '{pattern}' in {path}."}
        return {"ok": True, "path": path, "targets": targets,
                "msg": (f"Plan: move {len(targets)} file(s) matching '{pattern}' to a recovery folder "
                        f"(not permanently deleted):\n  " + "\n  ".join(targets[:20]) +
                        (f"\n  ...and {len(targets)-20} more" if len(targets) > 20 else ""))}

    def execute_delete(self, plan: dict) -> str:
        path = plan["path"]
        trash = os.path.join(path, RECOVERY_DIR, datetime.now().strftime("%Y%m%d_%H%M%S"))
        os.makedirs(trash, exist_ok=True)
        moved = 0
        for name in plan["targets"]:
            try:
                shutil.move(os.path.join(path, name), os.path.join(trash, name))
                moved += 1
            except OSError:
                pass
        return (f"Moved {moved} file(s) to {trash}. "
                "They're recoverable there — nothing was permanently deleted.")
